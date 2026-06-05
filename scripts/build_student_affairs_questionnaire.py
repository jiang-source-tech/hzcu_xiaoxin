from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "docs" / "学生事务知识库校准问题清单.docx"


SECTIONS = [
    (
        "一、可直接短答：地点、窗口、办公室",
        [
            "校园卡丢了去哪里补办？请只说地点。",
            "校园卡服务中心现在在哪里？请只说地点。",
            "信电学院学工办办公室在哪里？请只说地点。",
            "信电学院教学办办公室在哪里？请只说地点。",
            "信电学院党政办办公室在哪里？请只说地点。",
            "行政楼自助打印终端在哪里？请只说地点。",
            "校医务室在哪里？请只说地点。",
            "心理咨询室在哪里？请只说地点。",
            "学生事务办公室在哪里？请只说地点。",
            "团务问题一般去哪里问？请只说办公室或渠道。",
            "入党材料或党员发展问题一般去哪里问？请只说办公室或渠道。",
            "就业材料或毕业去向问题一般去哪里问？请只说办公室或渠道。",
        ],
    ),
    (
        "二、可直接短答：电话、入口、小程序",
        [
            "校医务室电话是多少？请只说电话。",
            "心理咨询预约电话是多少？请只说电话。",
            "学校心理咨询一般通过什么公众号或渠道预约？请一句话回答。",
            "校园卡补办是否需要通过什么线上入口？如果没有，请说线下地点。",
            "寝室报修一般在哪个系统或入口办理？请一句话回答。",
            "桶装水订购一般在哪个系统或入口办理？请一句话回答。",
            "生活缴费一般在哪个系统或入口办理？请一句话回答。",
            "学生财务银行卡绑定一般从哪个系统进入？请一句话回答。",
            "档案远程查档一般通过什么公众号或入口？请一句话回答。",
            "团组织关系转接一般在哪个平台办理？请一句话回答。",
        ],
    ),
    (
        "三、适合小芯回答的简短流程",
        [
            "学生证遗失补办，一般第一步找谁或去哪里？请一句话回答。",
            "火车票学生优惠卡办理，一般第一步找谁或去哪里？请一句话回答。",
            "火车票学生优惠区间修改，一般找谁处理？请一句话回答。",
            "寝室怎么报修？请只说入口或第一步。",
            "寒暑假留校住宿一般在哪里申请？请只说入口或第一步。",
            "收到公寓违规推送后，如果要申诉，一般在哪里操作？请一句话回答。",
            "怎么申请无障碍寝室？请只说第一步找谁。",
            "想咨询就业协议怎么签，一般先问谁？请一句话回答。",
        ],
    ),
    (
        "四、小芯不展开政策：只确认该找谁问",
        [
            "学费缴纳、退费、财务问题，学生一般应该问谁确认？",
            "选课、退课、补考报名这类教务流程，学生一般应该问谁确认？",
            "转专业政策或申请条件，学生一般应该问谁确认？",
            "医保参保、报销比例、异地备案这类政策，学生一般应该问谁确认？",
            "奖学金、助学金、困难认定这类资助政策，学生一般应该问谁确认？",
            "助学贷款材料和申请条件，学生一般应该问谁确认？",
            "入伍优待政策、报名条件这类政策，学生一般应该问谁确认？",
            "党员发展条件和流程，学生一般应该问谁确认？",
            "毕业生档案派遣、就业协议这类毕业事务，学生一般应该问谁确认？",
            "寝室调换、床位调整这类住宿审批，学生一般应该问谁确认？",
        ],
    ),
    (
        "五、容易过期：只收短事实，不收长解释",
        [
            "学生公寓自习室开放时间现在是多少？请只说时间。",
            "校医务室工作时间是多少？如果不确定，请说问校医务室或辅导员。",
            "心理咨询预约时间一般是什么时候？如果不确定，请说问心理中心或辅导员。",
            "校园卡服务中心工作时间是多少？如果不确定，请说问服务中心。",
            "学生事务办公室工作时间是多少？如果不确定，请说问辅导员。",
            "学校哪些事务适合问辅导员，而不是问智能助手？请列 3 个例子即可。",
            "如果学生问具体政策金额、条件、审批结果，智能助手应该怎么提醒？请一句话回答。",
        ],
    ),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = "DADCE0", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=140, bottom=120, end=140) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_font(run, size: int | float | None = None, bold=False, color=None) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc: Document, text: str = "", style: str | None = None, **run_style):
    paragraph = doc.add_paragraph(style=style)
    if text:
        run = paragraph.add_run(text)
        set_font(run, **run_style)
    return paragraph


def add_response_box(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(6.45)
    cell = table.cell(0, 0)
    cell.width = Inches(6.45)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_shading(cell, "F8FAFC")
    set_cell_borders(cell)
    set_cell_margins(cell)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("官方助手回答：")
    set_font(r, 10.5, bold=True, color="0B2545")

    for _ in range(4):
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(" ")
        set_font(r, 10.5)

    meta = cell.add_paragraph()
    meta.paragraph_format.space_before = Pt(2)
    meta.paragraph_format.space_after = Pt(0)
    r = meta.add_run("回答时间：__________    与现有知识库冲突：□ 无  □ 有    备注：________________")
    set_font(r, 9.5, color="555555")

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Heading 1", 18, "0B2545", 16, 8),
        ("Heading 2", 14, "1F4D78", 12, 6),
        ("Heading 3", 11.5, "0B2545", 8, 3),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_footer(section, text: str) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_font(run, 9, color="666666")


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    add_footer(section, "学生事务知识库校准问题清单")

    configure_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("学生事务知识库校准问题清单")
    set_font(run, 24, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("开发板/TTS 适配版：优先收集地点、窗口、入口、电话和“该问谁”，避免沉淀长政策原文。")
    set_font(run, 11, color="555555")

    note = doc.add_table(rows=1, cols=1)
    note.alignment = WD_TABLE_ALIGNMENT.LEFT
    note.autofit = False
    note.allow_autofit = False
    note.columns[0].width = Inches(6.45)
    cell = note.cell(0, 0)
    cell.width = Inches(6.45)
    set_cell_shading(cell, "EEF4FB")
    set_cell_borders(cell, color="B8C7D9")
    set_cell_margins(cell, top=120, bottom=120)
    p = cell.paragraphs[0]
    r = p.add_run("使用方法：")
    set_font(r, 10.5, bold=True, color="0B2545")
    r = p.add_run("把每个问题复制到学校智能助手里，把得到的回答粘贴到对应答案区。回答尽量控制在一两句话；政策、金额、条件、审批类内容只记录“建议问辅导员/相关老师”，不要粘贴长篇解释。")
    set_font(r, 10.5, color="0B2545")

    question_no = 1
    for section_title, questions in SECTIONS:
        doc.add_heading(section_title, level=1)
        for question in questions:
            q = doc.add_paragraph(style="Heading 3")
            r = q.add_run(f"{question_no}. {question}")
            set_font(r, 11.5, bold=True, color="0B2545")
            add_response_box(doc)
            question_no += 1

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_heading("补充问题", level=1)
    add_paragraph(
        doc,
        "如果学校智能助手在追问中给出了新的高频问题，可以继续记录在这里。",
        size=10.5,
        color="555555",
    )
    for i in range(1, 7):
        q = doc.add_paragraph(style="Heading 3")
        r = q.add_run(f"补充 {i}. 问题：")
        set_font(r, 11.5, bold=True, color="0B2545")
        add_response_box(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
